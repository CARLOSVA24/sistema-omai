
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_profound_justification():
    doc = docx.Document()

    # Style configuration
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # Header con mayor peso institucional
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run("ARMADA DEL ECUADOR\n")
    run.bold = True
    run.font.size = Pt(14)
    run = header.add_run("GRUPO DE TAREA 100.51 \"GUAYAS\"\n")
    run.bold = True
    run.font.size = Pt(12)
    run = header.add_run("PUESTO DE MANDO CENTRAL - SISTEMA OMAI")
    run.bold = True
    run.font.size = Pt(11)

    doc.add_paragraph("\n")

    # Title - Extremadamente convincente
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("INFORME TÉCNICO DE NECESIDAD IMPERIOSA Y SEGURIDAD OPERACIONAL\n")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(185, 28, 28) # Rojo oscuro para urgencia
    
    run = title.add_run("TRANSICIÓN DEL MODELO ANALÓGICO DISPERSO AL SISTEMA DE MANDO Y CONTROL INTEGRADO (OMAI)")
    run.bold = True
    run.font.size = Pt(12)
    run.underline = True

    doc.add_paragraph("\n")

    # Metadata
    def add_meta(label, value):
        p = doc.add_paragraph()
        run = p.add_run(f"{label}: ")
        run.bold = True
        p.add_run(value)

    add_meta("FECHA", "02 de junio de 2026")
    add_meta("DE", "EQUIPO DE DESARROLLO Y TÁCTICAS DIGITALES GT 100.51")
    add_meta("PARA", "COMANDANTE DEL GRUPO DE TAREA 100.51")
    add_meta("ASUNTO", "JUSTIFICACIÓN ESTRATÉGICA PARA LA IMPLEMENTACIÓN INMEDIATA DEL SISTEMA OMAI")

    doc.add_paragraph("-" * 80)

    # 1. DIAGNÓSTICO DE VULNERABILIDAD CRÍTICA
    h = doc.add_paragraph()
    run = h.add_run("1. DIAGNÓSTICO DE VULNERABILIDAD CRÍTICA")
    run.bold = True
    run.font.size = Pt(13)
    
    p = doc.add_paragraph(
        "Tras un análisis exhaustivo de los flujos de información en el Puesto de Mando del GT 100.51, se ha detectado una "
        "VULNERABILIDAD CRÍTICA sistémica. La infraestructura actual basada en hojas de cálculo (Excel) y repositorios en la "
        "nube (Google Drive) ha dejado de ser una herramienta de apoyo para convertirse en un riesgo para la seguridad y "
        "un cuello de botella para la efectividad operacional."
    )
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # 2. LA FALACIA DE LA SEGURIDAD EN LA NUBE PÚBLICA (DRIVES)
    h = doc.add_paragraph()
    run = h.add_run("2. LA INSUFICIENCIA DE LA SEGURIDAD EN SOPORTES TRADICIONALES")
    run.bold = True
    run.font.size = Pt(13)

    p = doc.add_paragraph(
        "• SOBERANÍA DE DATOS: Actualmente, información sensible (nombres, cédulas, capacidades operativas y ubicaciones "
        "de personal militar) reside en servidores de terceros fuera del control soberano de la Armada. Esto representa "
        "una exposición intolerable ante posibles ataques de ingeniería social o filtraciones.\n"
        "• FRAGMENTACIÓN TÁCTICA: La dispersión en múltiples archivos de Excel impide la visión periférica del Comandante. "
        "No es posible tomar decisiones en tiempo real si los datos deben ser consolidados manualmente en momentos de crisis.\n"
        "• INTEGRIDAD COMPROMETIDA: Los sistemas actuales carecen de validación. Cualquier usuario con acceso puede alterar, "
        "borrar o corromper registros históricos, sin posibilidad de auditoría o recuperación inmediata."
    )
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # 3. EL IMPACTO EN LA EFECTIVIDAD OPERATIVA
    h = doc.add_paragraph()
    run = h.add_run("3. IMPACTO EN LA CAPACIDAD DE MANDO, CONTROL Y RESPUESTA")
    run.bold = True
    run.font.size = Pt(13)

    p = doc.add_paragraph(
        "La ineficiencia administrativa es, en última instancia, una debilidad operativa. El tiempo que el personal destina a "
        "la corrección de errores en Excels y la búsqueda de información dispersa es tiempo restado a la planificación estratégica "
        "y a la reacción táctica. En el contexto de seguridad actual, un retraso de 30 minutos en la consolidación de un parte "
        "puede significar la pérdida de una ventaja táctica sobre las amenazas."
    )
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # 4. PROPUESTA DE VALOR: SISTEMA OMAI
    h = doc.add_paragraph()
    run = h.add_run("4. EL SISTEMA OMAI COMO SOLUCIÓN ESTRATÉGICA")
    run.bold = True
    run.font.size = Pt(13)

    p = doc.add_paragraph(
        "El Sistema OMAI no es una simple base de datos; es un multiplicador de fuerzas diseñado para:\n"
        "• CENTRALIZACIÓN TOTAL: Un solo repositorio blindado para Personal, Logística, Operaciones e Inteligencia.\n"
        "• BLINDAJE DE ACCESO (RBAC): Control absoluto de quién accede a qué información según su rol y necesidad de conocer.\n"
        "• AUTOMATIZACIÓN MILITAR: Generación de Órdenes de Patrulla y Partes en segundos, con estandarización institucional.\n"
        "• INTELIGENCIA GEOPESCUENCIAL: Visualización dinámica en mapas para detectar patrullajes superpuestos o zonas desatendidas."
    )
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # 5. CONCLUSIÓN Y DICTAMEN DE PUESTA EN VIGOR
    h = doc.add_paragraph()
    run = h.add_run("5. CONCLUSIÓN Y DICTAMEN")
    run.bold = True
    run.font.size = Pt(13)

    p = doc.add_paragraph(
        "La superioridad informativa es la base de la superioridad táctica. Continuar operando bajo un modelo de "
        "información dispersa es una negligencia que el GT 100.51 no puede permitirse. La implementación del SISTEMA OMAI "
        "es el paso obligatorio para profesionalizar la gestión del Puesto de Mando y garantizar que la información trabaje "
        "a favor de la misión y no en contra del personal.\n\n"
        "Se recomienda la PUESTA EN VIGOR INMEDIATA del programa para salvaguardar la seguridad de la información "
        "y optimizar la respuesta operativa institucional."
    )
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_paragraph("\n\n")
    
    # Signature
    sig = doc.add_paragraph()
    sig.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sig.add_run("-" * 50 + "\n")
    run = sig.add_run("PERSONAL TÉCNICO DE MANDO Y CONTROL\nSISTEMA OMAI - GT 100.51")
    run.bold = True

    doc.save("INFORME_JUSTIFICACION_ESTRATEGICA_PROFUNDO.docx")
    print("Archivo INFORME_JUSTIFICACION_ESTRATEGICA_PROFUNDO.docx generado con éxito.")

if __name__ == "__main__":
    create_profound_justification()
