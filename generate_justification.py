
import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_justification():
    doc = docx.Document()

    # Style configuration
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # Header
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run("ARMADA DEL ECUADOR\nGRUPO DE TAREA 100.51\nSISTEMA OMAI")
    run.bold = True
    run.font.size = Pt(12)

    doc.add_paragraph("\n")

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("INFORME TÉCNICO DE JUSTIFICACIÓN\nDigitalización y Centralización de Información Operativa")
    run.bold = True
    run.font.size = Pt(14)
    run.underline = True

    doc.add_paragraph("\n")

    # Metadata
    p = doc.add_paragraph()
    p.add_run("FECHA: ").bold = True
    p.add_run("02 de junio de 2026")
    
    p = doc.add_paragraph()
    p.add_run("PARA: ").bold = True
    p.add_run("Mando del GT 100.51 / Autoridades Competentes")
    
    p = doc.add_paragraph()
    p.add_run("DE: ").bold = True
    p.add_run("Desarrollo de Sistemas OMAI")
    
    p = doc.add_paragraph()
    p.add_run("ASUNTO: ").bold = True
    p.add_run("Justificación para la puesta en vigor del sistema de gestión centralizada.")

    doc.add_paragraph("-" * 80)

    # Content Sections
    sections = [
        ("1. ANTECEDENTES", 
         "El Puesto de Mando del Grupo de Tarea 100.51 visualiza la necesidad imperativa de modernizar sus procesos de gestión de información. La naturaleza de las operaciones militares de seguridad ciudadana requiere agilidad, precisión y, sobre todo, integridad de los datos."),
        
        ("2. PROBLEMÁTICA IDENTIFICADA", 
         "Actualmente, la gestión de la información operativa (Órdenes de Patrulla, Registro de Personal, Partes al Instante) se realiza mediante herramientas ofimáticas tradicionales como Microsoft Excel y almacenamiento en la nube (Google Drive). Esta metodología presenta los siguientes inconvenientes:\n"
         "• Dispersión de la información en múltiples archivos y carpetas.\n"
         "• Dificultad para consolidar estadísticas en tiempo real.\n"
         "• Duplicidad de registros y errores humanos en la transcripción.\n"
         "• Dependencia de la conectividad externa para acceder a archivos compartidos."),
        
        ("3. RIESGOS DE SEGURIDAD DE LA INFORMACIÓN", 
         "El uso de archivos Excel y Drives no permite salvaguardar la seguridad institucional por las siguientes razones:\n"
         "• Falta de Control de Acceso Granular: No se puede restringir quién ve o edita secciones específicas de un archivo de manera eficiente.\n"
         "• Vulnerabilidad ante Eliminación Accidental: Un usuario puede borrar celdas o archivos completos sin dejar rastro de auditoría.\n"
         "• Riesgo de Filtración: La información militar sensible reside en servidores externos de terceros, fuera de la infraestructura controlada por la institución.\n"
         "• Integridad Comprometida: No existe validación de datos, lo que permite el ingreso de información inconsistente o incompleta."),
        
        ("4. LA SOLUCIÓN: SISTEMA OMAI", 
         "Se ha desarrollado un sistema web centralizado que utiliza una base de datos relacional (SQLite) y una interfaz intuitiva con las siguientes características:\n"
         "• Repositorio Único: Toda la información de personal, logística e inteligencia se almacena en un solo punto.\n"
         "• Seguridad por Roles (RBAC): El acceso es validado mediante credenciales y roles (Administrador, Jefe, Personal, etc.), limitando las funciones según el rango.\n"
         "• Automatización Documental: Generación instantánea de Órdenes de Patrulla y Reportes en formato PDF con estándares militares.\n"
         "• Visualización Geográfica: Mapa interactivo para el registro y análisis de incidentes en tiempo real."),
        
        ("5. BENEFICIOS ESPERADOS", 
         "• Optimización del tiempo en la elaboración de partes y órdenes (reducción del 70% en tiempo administrativo).\n"
         "• Precisión estadística para la toma de decisiones del Comandante del GT 100.51.\n"
         "• Centralización del control de personal operativo y administrativo.\n"
         "• Garantía de la persistencia histórica de los datos para análisis posteriores."),
        
        ("6. CONCLUSIÓN", 
         "La implementación y puesta en vigor del Programa OMAI no es solo una mejora tecnológica, sino una necesidad estratégica para garantizar la superioridad en la gestión de información y la seguridad de las operaciones en la jurisdicción del GT 100.51.")
    ]

    for title_text, body_text in sections:
        h = doc.add_paragraph()
        run = h.add_run(title_text)
        run.bold = True
        run.font.size = Pt(12)
        
        p = doc.add_paragraph(body_text)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_paragraph("\n\n")
    
    # Signature
    sig = doc.add_paragraph()
    sig.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sig.add_run("-" * 40 + "\n")
    sig.add_run("EQUIPO DE DESARROLLO OMAI\nSOPORTE TÉCNICO GT 100.51")

    doc.save("INFORME_JUSTIFICACION_SISTEMA.docx")
    print("Archivo INFORME_JUSTIFICACION_SISTEMA.docx generado con éxito.")

if __name__ == "__main__":
    create_justification()
