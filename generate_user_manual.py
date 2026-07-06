
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_manual():
    doc = docx.Document()

    # Base Styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Inter'
    font.size = Pt(10)

    # Helper function for headings
    def add_custom_heading(text, level, bold=True, size=14, color=None):
        h = doc.add_paragraph()
        if level > 0:
            h.paragraph_format.space_before = Pt(20)
        run = h.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = color
        return h

    # Portada
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run("ARMADA DEL ECUADOR\nGRUPO DE TAREA 100.51")
    run.bold = True
    run.font.size = Pt(16)

    doc.add_paragraph("\n" * 5)
    
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("MANUAL DEL USUARIO\nSISTEMA OMAI")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(37, 99, 235) # Blue

    doc.add_paragraph("\n" * 2)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("PLATAFORMA CENTRALIZADA DE MANDO Y CONTROL")
    run.italic = True
    run.font.size = Pt(14)

    doc.add_paragraph("\n" * 10)
    version = doc.add_paragraph()
    version.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    version.add_run("Versión 2.0\nJunio 2026")

    doc.add_page_break()

    # --- TABLA DE CONTENIDO ---
    add_custom_heading("1. INTRODUCCIÓN", 1, size=16, color=RGBColor(30, 58, 138))
    doc.add_paragraph("El Sistema OMAI (Operaciones Militares de Ámbito Interno) es la herramienta tecnológica oficial del GT 100.51 para la gestión de personal, logística, inteligencia y planificación operativa. Este manual proporciona las instrucciones detalladas para el uso correcto de la plataforma.")

    # --- ACCESO ---
    add_custom_heading("2. INICIO Y CONTROL DE ACCESO", 1, size=16, color=RGBColor(30, 58, 138))
    p = doc.add_paragraph("El sistema requiere una autenticación previa para garantizar que solo personal autorizado acceda a la información sensible.")
    
    steps = [
        "Ejecute el archivo 'INICIAR_PROGRAMA.bat' en el equipo servidor.",
        "Abra el navegador Chrome e ingrese a http://localhost:3000.",
        "En la pantalla de Login, seleccione su ROL (EJ: PERSONAL OMAI, JEFE OMAI).",
        "Ingrese su contraseña institucional.",
        "Haga clic en 'INGRESAR AL SISTEMA'."
    ]
    for step in steps:
        doc.add_paragraph(step, style='List Number')

    # --- MODULO PERSONAL ---
    add_custom_heading("3. MÓDULO DE PERSONAL", 1, size=16, color=RGBColor(30, 58, 138))
    doc.add_paragraph("Permite la administración completa del talento humano asignado al GT 100.51.")
    
    add_custom_heading("3.1 Registro de Personal", 2, size=14)
    doc.add_paragraph("Para dar el alta a un nuevo elemento:\n1. Ingrese los datos básicos (Grado, Especialidad, Nombres).\n2. Importante: Ingrese la Cédula (10 dígitos) para evitar duplicados.\n3. Seleccione el Grupo de Rotación.\n4. Haga clic en Guardar.")
    
    add_custom_heading("3.2 Estadísticas y Dashboard", 2, size=14)
    doc.add_paragraph("El sistema genera automáticamente gráficos de:\n• Efectivos por Puesto.\n• Distribución por Repartos.\n• Condición Operativa (Operativo, Permiso, Franco).")

    # --- MODULO OPERACIONES ---
    add_custom_heading("4. MÓDULO DE OPERACIONES", 1, size=16, color=RGBColor(30, 58, 138))
    
    add_custom_heading("4.1 Planificación Diaria", 2, size=14)
    doc.add_paragraph("Sección crítica para el Puesto de Mando. Permite registrar las tareas planificadas para el día, asignando distritos y responsables.")
    
    add_custom_heading("4.2 Generación de Órdenes de Patrulla (ORDPAT)", 2, size=14)
    doc.add_paragraph("El sistema permite generar documentos militares estandarizados:\n1. Seleccione 'Orden de Patrulla GT ECHO' o 'GT 100.51'.\n2. Edite los campos de Misión, Ejecución y Tareas.\n3. El Anexo 'A' se autocompleta con el personal seleccionado.\n4. Use el botón 'Imprimir' para generar el PDF con el formato reglamentario.")

    add_custom_heading("4.3 Partes al Instante", 2, size=14)
    doc.add_paragraph("Para comunicar novedades en tiempo real:\n1. Llene el formulario con el tipo de novedad.\n2. Ingrese coordenadas o seleccione en el mapa.\n3. Adjunte fotos del evento.\n4. El reporte se guardará en el histórico y podrá ser exportado.")

    # --- MODULO INTELIGENCIA ---
    add_custom_heading("5. MÓDULO DE INTELIGENCIA Y MAPA", 1, size=16, color=RGBColor(30, 58, 138))
    doc.add_paragraph("El mapa interactivo es el centro de análisis geoespacial.")
    
    add_custom_heading("5.1 Registro de Incidentes", 2, size=14)
    doc.add_paragraph("1. Haga clic en el punto exacto del mapa donde ocurrió el delito.\n2. Las coordenadas se cargarán automáticamente.\n3. Seleccione el tipo de delito (Robo, Sicariato, etc.).\n4. Guarde el registro para que aparezca como un marcador en el mapa.")
    
    add_custom_heading("5.2 Herramientas de Dibujo", 2, size=14)
    doc.add_paragraph("Use la barra superior para dibujar Polígonos de sectores de responsabilidad o medir distancias con la Regla.")

    # --- LOGISTICA ---
    add_custom_heading("6. MÓDULO DE LOGÍSTICA", 1, size=16, color=RGBColor(30, 58, 138))
    doc.add_paragraph("Control de vehículos terrestres. Registre placas, modelo, estado mecánico y conductor asignado.")

    # --- SEGURIDAD ---
    add_custom_heading("7. SEGURIDAD Y RESPALDOS", 1, size=16, color=RGBColor(30, 58, 138))
    doc.add_paragraph("• Cambio de Contraseña: Solo el Administrador puede gestionar credenciales en el menú inferior.\n• Errores: En caso de mal funcionamiento, verifique que la consola del Servidor esté abierta.")

    doc.add_page_break()
    doc.add_paragraph("Fin del Documento", style='Intense Quote')

    doc.save("MANUAL_USUARIO_DETALLADO_OMAI.docx")
    print("Archivo MANUAL_USUARIO_DETALLADO_OMAI.docx generado con éxito.")

# Dummy function to mimic a variable size
def size(s): return s

if __name__ == "__main__":
    create_manual()
