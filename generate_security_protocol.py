
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_security_protocol():
    doc = docx.Document()

    # Style configuration
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)

    def add_h(text, size=14, color=None):
        h = doc.add_paragraph()
        run = h.add_run(text)
        run.bold = True
        run.font.size = Pt(size)
        if color: run.font.color.rgb = color
        return h

    # Header
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run("PROTOCOLO DE INTEGRIDAD Y BLINDAJE DEL SISTEMA\n")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(127, 29, 29) # Rojo oscuro

    doc.add_paragraph("\n")
    doc.add_paragraph(
        "Para garantizar que el código fuente no sea alterado ni manipulado al momento de su despliegue "
        "en diferentes estaciones de trabajo, se deben seguir los siguientes pasos técnicos de protección de integridad."
    )

    add_h("1. LIMPIEZA DEL PAQUETE DE DESPLIEGUE", 13, RGBColor(31, 41, 55))
    doc.add_paragraph(
        "Antes de copiar la carpeta a un nuevo ordenador, elimine los archivos que permiten la edición. "
        "Deje solamente lo necesario para la ejecución:\n"
        "• ELIMINAR: Todos los archivos con extensión .py (scripts de desarrollo).\n"
        "• ELIMINAR: Bases de datos temporales (test.db) si existen.\n"
        "• ELIMINAR: Manuales en formato editable (.md) si ya tiene los Word.\n"
        "• MANTENER: SISTEMA_OMAI.exe, index.html, style.css, script_v2.js, la carpeta node_modules y database.sqlite."
    )

    add_h("2. PROTECCIÓN DE ARCHIVOS (ATRIBUTO SOLO LECTURA)", 13, RGBColor(31, 41, 55))
    doc.add_paragraph(
        "Puede 'congelar' los archivos usando el comando de sistema. Ejecute esto en una terminal dentro de la carpeta:\n"
        "Comando: attrib +r +s +h *.js\n"
        "Esto hará que los archivos JavaScript sean de solo lectura, de sistema y ocultos, impidiendo su edición accidental o malintencionada."
    )

    add_h("3. PERMISOS DE SEGURIDAD NTFS (EL MÁS EFECTIVO)", 13, RGBColor(31, 41, 55))
    doc.add_paragraph(
        "Una vez copiada la carpeta al disco C:\ del nuevo ordenador:\n"
        "1. Clic derecho sobre la carpeta del sistema > Propiedades.\n"
        "2. Pestaña 'Seguridad' > botón 'Opciones avanzadas'.\n"
        "3. Deshabilite la herencia (Diga 'Convertir permisos heredados').\n"
        "4. Al usuario 'Usuarios' o 'Todos', cámbiele los permisos para que solo tenga:\n"
        "   - Lectura y ejecución.\n"
        "   - Mostrar el contenido de la carpeta.\n"
        "   - Lectura.\n"
        "5. Quite los permisos de 'Escritura' y 'Modificación'.\n"
        "   Nota: El archivo 'database.sqlite' SÍ debe tener permisos de escritura para que el programa guarde datos."
    )

    add_h("4. USO DE ACCESOS DIRECTOS", 13, RGBColor(31, 41, 55))
    doc.add_paragraph(
        "Nunca dé acceso directo a la carpeta del programa a los usuarios finales. Cree un acceso directo del "
        "archivo 'INICIAR_PROGRAMA.bat' en el escritorio y cámbiele el ícono. Los usuarios no necesitan ver "
        "los archivos internos para operar el sistema."
    )

    add_h("5. INTEGRIDAD DE LA BASE DE DATOS", 13, RGBColor(31, 41, 55))
    doc.add_paragraph(
        "Para evitar modificaciones directas en la base de datos sin usar la interfaz del sistema, configure "
        "una contraseña en la gestión de roles del programa para que solo el rol 'ADMINISTRADOR' tenga acceso "
        "a las funciones críticas."
    )

    doc.add_paragraph("\n\n")
    sig = doc.add_paragraph("Certificación de Integridad de Software\nEquipo de Desarrollo OMAI")
    sig.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save("PROTOCOLO_INTEGRIDAD_SISTEMA.docx")
    print("Archivo PROTOCOLO_INTEGRIDAD_SISTEMA.docx generado con éxito.")

if __name__ == "__main__":
    create_security_protocol()
