
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_blindaje_guide():
    doc = docx.Document()

    # Estilos Globales
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    def add_title(text, size=18, color=RGBColor(31, 73, 125)):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(size)
        run.font.color.rgb = color
        return p

    def add_heading(text, level=1):
        h = doc.add_heading(text, level)
        return h

    # --- PORTADA ---
    add_title("PROTOCOLO DE INSTALACIÓN Y BLINDAJE DEL SISTEMA", 22)
    add_title("SISTEMA OMAI GT 100.51", 16, RGBColor(128, 128, 128))
    
    doc.add_paragraph("\n" * 5)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("FINALIDAD: Garantizar la integridad del código fuente y asegurar la correcta puesta en vigor en nuevos terminales.")
    run.font.italic = True
    
    doc.add_paragraph("\n" * 8)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("DEPARTAMENTO DE INFORMÁTICA / SOPORTE TÉCNICO\nGT 100.51 - PUESTO DE MANDO")

    doc.add_page_break()

    # --- ARCHIVOS NECESARIOS ---
    add_heading("1. ARCHIVOS ESENCIALES PARA LA COPIA", 1)
    doc.add_paragraph("Para que el sistema funcione en otro ordenador, DEBEN copiarse los siguientes archivos en una sola carpeta (ejemplo C:\\SISTEMA_OMAI):")
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Archivo / Carpeta'
    hdr_cells[1].text = 'Función Crítica'

    files = [
        ("index.html", "Estructura visual y menús del sistema."),
        ("style.css", "Diseño, colores y animaciones."),
        ("database.sqlite", "Corazón del sistema: contiene todos los registros."),
        ("INICIAR_PROGRAMA.bat", "Lanzador que activa el servidor y abre el navegador."),
        ("ESCUDOARMADA.jpg", "Identidad institucional visual."),
        ("package.json", "Configuración técnica del servidor Node.js."),
        ("SOPORTE_TECNICO/", "Cerebro del programa (server.js y script_v2.js)."),
        ("node_modules/", "Librerías necesarias para el funcionamiento (NO BORRAR).")
    ]

    for item, desc in files:
        row_cells = table.add_row().cells
        row_cells[0].text = item
        row_cells[1].text = desc

    # --- HERRAMIENTA AUTOMÁTICA ---
    add_heading("2. HERRAMIENTA DE DESPLIEGUE AUTOMÁTICO", 1)
    doc.add_paragraph("El sistema incluye una utilidad llamada HERRAMIENTA_DESPLIEGUE_ADMIN.bat. Se recomienda su uso para generar un paquete 'blindado' automáticamente:")
    
    doc.add_paragraph("1. Ejecute el archivo en el ordenador de origen.", style='List Number')
    doc.add_paragraph("2. Ingrese la clave de administrador (admin).", style='List Number')
    doc.add_paragraph("3. El sistema creará una carpeta en C:\\OMAI_DISTRIBUCION_LISTA con solo lo necesario.", style='List Number')
    doc.add_paragraph("4. Copie esa carpeta generada al nuevo ordenador.", style='List Number')

    # --- BLINDAJE Y SEGURIDAD ---
    add_heading("3. PROTOCOLO DE BLINDAJE (PROTECCIÓN ANT-MODIFICACIÓN)", 1)
    doc.add_paragraph("Para asegurar que el código fuente no sea alterado en el nuevo equipo, siga estas instrucciones post-instalación:")
    
    p = doc.add_paragraph()
    run = p.add_run("A. Atributos de Solo Lectura (Básico):")
    run.bold = True
    doc.add_paragraph("Una vez copiado al nuevo PC, abra una consola (CMD) en la carpeta del sistema y ejecute:")
    p = doc.add_paragraph("attrib +r *.html\nattrib +r *.css\nattrib +r SOPORTE_TECNICO\\*.js", style='No Spacing')
    p.paragraph_format.left_indent = Inches(0.5)

    p = doc.add_paragraph()
    run = p.add_run("B. Permisos NTFS (Avanzado/Recomendado):")
    run.bold = True
    doc.add_paragraph("1. Clic derecho sobre la carpeta del sistema > Propiedades > Seguridad.")
    doc.add_paragraph("2. Seleccione el grupo 'Usuarios' o el nombre del operador local.")
    doc.add_paragraph("3. Haga clic en Editar y DENIEGUE el permiso de 'Escritura' y 'Modificación'.")
    doc.add_paragraph("4. Deje permitidos únicamente 'Lectura' y 'Ejecución'.")
    
    doc.add_paragraph("Con esto, el operador podrá usar el programa pero recibirá un error de Windows si intenta borrar o cambiar una sola línea de código.")

    # --- PUESTA EN VIGOR ---
    add_heading("4. PUESTA EN VIGOR EN RED LOCAL", 1)
    doc.add_paragraph("Si el nuevo ordenador funcionará como servidor para otros terminales:")
    doc.add_paragraph("• IP ESTÁTICA: Configure una IP fija en el ordenador para que el acceso no cambie.")
    doc.add_paragraph("• FIREWALL: Abra el puerto 3000 en el Firewall de Windows para permitir la entrada de otros equipos.")
    doc.add_paragraph("• ACCESO: Los otros equipos entrarán mediante el navegador digitando: http://[IP_DEL_SVR]:3000")

    doc.add_paragraph("\n" * 3)
    p = doc.add_paragraph("MANTENGA ESTA GUÍA EN UN LUGAR SEGURO")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save("MANUAL_INSTALACION_Y_BLINDAJE.docx")
    print("Manual generado: MANUAL_INSTALACION_Y_BLINDAJE.docx")

if __name__ == "__main__":
    create_blindaje_guide()
