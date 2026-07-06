
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_tech_guide():
    doc = docx.Document()

    # Base Styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Segoe UI' if 'Segoe UI' in [s.name for s in doc.styles] else 'Arial'
    font.size = Pt(10)

    def add_heading(text, level, size=14, color=None):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(15)
        run = h.add_run(text)
        run.bold = True
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = color
        return h

    # Portada Técnica
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run("GUÍA TÉCNICA DE INSTALACIÓN Y DESPLIEGUE EN RED\n")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(30, 64, 175)

    doc.add_paragraph("\n" * 2)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("SISTEMA OMAI GT 100.51")
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph("\n" * 10)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("Departamento de Soporte Técnico\nGT 100.51 - Mando y Control")

    doc.add_page_break()

    # --- CONTENIDO ---
    add_heading("1. ARQUITECTURA DEL SISTEMA", 1, 16, RGBColor(30, 58, 138))
    doc.add_paragraph(
        "El Sistema OMAI opera bajo una arquitectura de Servidor-Cliente. "
        "El 'Servidor' es el computador donde reside la base de datos y el motor del programa, "
        "mientras que los 'Clientes' son otros computadores que acceden a través de la red local."
    )

    # --- SERVIDOR ---
    add_heading("2. CONFIGURACIÓN DEL SERVIDOR (PC PRINCIPAL)", 1, 16, RGBColor(30, 58, 138))
    
    add_heading("2.1 Preparación de IP Estática (Recomendado)", 2, 12)
    doc.add_paragraph(
        "Para evitar que la dirección de acceso cambie cada vez que el equipo se reinicie, se debe asignar una IP fija:\n"
        "1. Vaya a Panel de Control > Centro de Redes > Cambiar configuración del adaptador.\n"
        "2. Clic derecho sobre su conexión (Ethernet/Wi-Fi) > Propiedades.\n"
        "3. Seleccione 'Protocolo de Internet versión 4 (TCP/IPv4)' > Propiedades.\n"
        "4. Asigne una dirección IP (Ej: 192.168.1.100) y guarde los cambios."
    )

    add_heading("2.2 Instalación del Software", 2, 12)
    doc.add_paragraph(
        "1. Copie la carpeta 'registro de delitos' íntegramente al disco local C:\ del servidor.\n"
        "2. No instale Node.js manualmente; el archivo 'SISTEMA_OMAI.exe' ya contiene el entorno de ejecución necesario.\n"
        "3. Localice el archivo 'INICIAR_PROGRAMA.bat' y envíelo al Escritorio como acceso directo."
    )

    add_heading("2.3 Apertura del Firewall de Windows", 2, 12)
    doc.add_paragraph(
        "Este es el paso más importante para que otros computadores puedan ver el sistema:\n"
        "1. Busque 'Firewall de Windows Defender con seguridad avanzada' en el menú inicio.\n"
        "2. Haga clic en 'Reglas de entrada' > 'Nueva regla'.\n"
        "3. Seleccione 'Puerto' > Siguiente > 'TCP' y en Puertos locales específicos escriba: 3000.\n"
        "4. Seleccione 'Permitir la conexión' > Siguiente.\n"
        "5. Marque Dominio, Privado y Público > Siguiente.\n"
        "6. Nombre de la regla: 'SISTEMA OMAI - PUERTO 3000' y Finalice."
    )

    # --- CLIENTES ---
    add_heading("3. CONFIGURACIÓN DE EQUIPOS CLIENTES", 1, 16, RGBColor(30, 58, 138))
    doc.add_paragraph(
        "En los computadores de las otras oficinas u operadoras:\n"
        "1. Asegúrese de que estén conectados a la misma red (mismo switch o misma Wi-Fi).\n"
        "2. Abra Google Chrome.\n"
        "3. En la barra de direcciones, escriba la IP del servidor seguida de ':3000'.\n"
        "   Ejemplo: http://192.168.1.100:3000\n"
        "4. ¡Listo! El sistema cargará de inmediato sin necesidad de instalar nada en el cliente."
    )

    # --- PUESTA EN VIGOR ---
    add_heading("4. RUTINA DE PUESTA EN VIGOR DIARIA", 1, 16, RGBColor(30, 58, 138))
    doc.add_paragraph(
        "Para que el sistema sea efectivo, siga esta rutina de inicio:\n"
        "1. El operador del Servidor debe encender el equipo a primera hora (07:00 AM).\n"
        "2. Ejecutar el acceso directo 'INICIAR_PROGRAMA.bat'.\n"
        "3. Mantener la ventana de la consola abierta (se puede minimizar pero NO cerrar).\n"
        "4. El servidor debe permanecer encendido durante todo el turno de guardia 24/7."
    )

    # --- RESPALDOS ---
    add_heading("5. POLÍTICA DE RESPALDO DE INFORMACIÓN", 1, 16, RGBColor(30, 58, 138))
    doc.add_paragraph(
        "Dado que el sistema centraliza información crítica, es imperativo:\n"
        "• Cada lunes, copiar el archivo 'database.sqlite' a un dispositivo externo (pendrive) o un disco de red fuera del equipo servidor.\n"
        "• Este archivo es el corazón del programa; si el disco del servidor falla, este respaldo permitirá recuperar todo el historial en minutos en cualquier otro computador."
    )

    doc.add_page_break()
    p = doc.add_paragraph("Fin de la Guía de Instalación")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save("GUIA_INSTALACION_SERVIDOR_RED.docx")
    print("Archivo GUIA_INSTALACION_SERVIDOR_RED.docx generado con éxito.")

if __name__ == "__main__":
    create_tech_guide()
